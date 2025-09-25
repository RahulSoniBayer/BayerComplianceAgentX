"""
DOCX parsing module for extracting placeholders and filling generated content.
Handles Word document templates with placeholder extraction and content insertion.
"""

import logging
import re
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from utils.validators import ContentValidator

logger = logging.getLogger(__name__)


@dataclass
class Placeholder:
    """Represents a placeholder found in a DOCX template."""
    text: str
    placeholder_type: str  # 'table', 'section', 'inline'
    context_type: str  # 'table', 'section'
    position_in_document: int
    paragraph_index: int
    run_index: int
    start_pos: int
    end_pos: int
    metadata: Dict[str, Any] = None
    
    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


class DOCXParser:
    """Main DOCX parsing class for extracting placeholders and filling content."""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        # Common placeholder patterns
        self.placeholder_patterns = [
            r'{{([^}]+)}}',  # {{placeholder}}
            r'\[\[([^\]]+)\]\]',  # [[placeholder]]
            r'<([^>]+)>',  # <placeholder>
            r'%([^%]+)%',  # %placeholder%
        ]
    
    def extract_placeholders(self, docx_path: str) -> List[Placeholder]:
        """
        Extract all placeholders from a DOCX template.
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            List[Placeholder]: List of found placeholders
        """
        try:
            doc = Document(docx_path)
            placeholders = []
            position_counter = 0
            
            # Process paragraphs
            for para_idx, paragraph in enumerate(doc.paragraphs):
                para_placeholders = self._extract_placeholders_from_paragraph(
                    paragraph, para_idx, position_counter
                )
                placeholders.extend(para_placeholders)
                position_counter += len(para_placeholders)
            
            # Process tables
            for table_idx, table in enumerate(doc.tables):
                table_placeholders = self._extract_placeholders_from_table(
                    table, table_idx, position_counter
                )
                placeholders.extend(table_placeholders)
                position_counter += len(table_placeholders)
            
            self.logger.info(f"Extracted {len(placeholders)} placeholders from {docx_path}")
            return placeholders
            
        except Exception as e:
            self.logger.error(f"Error extracting placeholders from {docx_path}: {str(e)}")
            raise
    
    def _extract_placeholders_from_paragraph(self, paragraph, para_idx: int, position_start: int) -> List[Placeholder]:
        """
        Extract placeholders from a paragraph.
        
        Args:
            paragraph: DOCX paragraph object
            para_idx: Paragraph index
            position_start: Starting position counter
            
        Returns:
            List[Placeholder]: List of placeholders found in this paragraph
        """
        placeholders = []
        position_counter = position_start
        
        for run_idx, run in enumerate(paragraph.runs):
            text = run.text
            if not text:
                continue
            
            # Find all placeholder patterns in this run
            for pattern in self.placeholder_patterns:
                matches = list(re.finditer(pattern, text))
                for match in matches:
                    placeholder_text = match.group(1).strip()
                    
                    # Validate placeholder text
                    if not ContentValidator.validate_placeholder_text(placeholder_text):
                        continue
                    
                    # Determine placeholder type and context
                    placeholder_type, context_type = self._classify_placeholder(
                        placeholder_text, paragraph, para_idx
                    )
                    
                    placeholder = Placeholder(
                        text=placeholder_text,
                        placeholder_type=placeholder_type,
                        context_type=context_type,
                        position_in_document=position_counter,
                        paragraph_index=para_idx,
                        run_index=run_idx,
                        start_pos=match.start(),
                        end_pos=match.end(),
                        metadata={
                            "pattern": pattern,
                            "full_match": match.group(0),
                            "is_in_table": False
                        }
                    )
                    placeholders.append(placeholder)
                    position_counter += 1
        
        return placeholders
    
    def _extract_placeholders_from_table(self, table, table_idx: int, position_start: int) -> List[Placeholder]:
        """
        Extract placeholders from a table.
        
        Args:
            table: DOCX table object
            table_idx: Table index
            position_start: Starting position counter
            
        Returns:
            List[Placeholder]: List of placeholders found in this table
        """
        placeholders = []
        position_counter = position_start
        
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para_idx, paragraph in enumerate(cell.paragraphs):
                    for run_idx, run in enumerate(paragraph.runs):
                        text = run.text
                        if not text:
                            continue
                        
                        # Find all placeholder patterns in this run
                        for pattern in self.placeholder_patterns:
                            matches = list(re.finditer(pattern, text))
                            for match in matches:
                                placeholder_text = match.group(1).strip()
                                
                                # Validate placeholder text
                                if not ContentValidator.validate_placeholder_text(placeholder_text):
                                    continue
                                
                                # Table placeholders are typically table context
                                placeholder = Placeholder(
                                    text=placeholder_text,
                                    placeholder_type="table",
                                    context_type="table",
                                    position_in_document=position_counter,
                                    paragraph_index=para_idx,
                                    run_index=run_idx,
                                    start_pos=match.start(),
                                    end_pos=match.end(),
                                    metadata={
                                        "pattern": pattern,
                                        "full_match": match.group(0),
                                        "is_in_table": True,
                                        "table_index": table_idx,
                                        "row_index": row_idx,
                                        "cell_index": cell_idx
                                    }
                                )
                                placeholders.append(placeholder)
                                position_counter += 1
        
        return placeholders
    
    def _classify_placeholder(self, text: str, paragraph, para_idx: int) -> Tuple[str, str]:
        """
        Classify a placeholder based on its text and context.
        
        Args:
            text: Placeholder text
            paragraph: DOCX paragraph object
            para_idx: Paragraph index
            
        Returns:
            Tuple[str, str]: (placeholder_type, context_type)
        """
        text_lower = text.lower()
        
        # Check for table-related keywords
        table_keywords = ['table', 'row', 'column', 'cell', 'data', 'list', 'item']
        if any(keyword in text_lower for keyword in table_keywords):
            return "table", "table"
        
        # Check for section-related keywords
        section_keywords = ['section', 'paragraph', 'description', 'overview', 'summary', 'details']
        if any(keyword in text_lower for keyword in section_keywords):
            return "section", "section"
        
        # Check paragraph style for context clues
        if paragraph.style and paragraph.style.name:
            style_lower = paragraph.style.name.lower()
            if 'heading' in style_lower or 'title' in style_lower:
                return "section", "section"
            if 'table' in style_lower:
                return "table", "table"
        
        # Default classification based on text length
        if len(text) > 50:
            return "section", "section"
        else:
            return "inline", "section"
    
    def fill_placeholders(self, docx_path: str, placeholder_fills: Dict[str, str], output_path: str) -> bool:
        """
        Fill placeholders in a DOCX template with generated content.
        
        Args:
            docx_path: Path to the original DOCX template
            placeholder_fills: Dictionary mapping placeholder text to filled content
            output_path: Path for the output filled document
            
        Returns:
            bool: True if successful
        """
        try:
            doc = Document(docx_path)
            filled_count = 0
            
            # Fill placeholders in paragraphs
            for paragraph in doc.paragraphs:
                filled_count += self._fill_placeholders_in_paragraph(paragraph, placeholder_fills)
            
            # Fill placeholders in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            filled_count += self._fill_placeholders_in_paragraph(paragraph, placeholder_fills)
            
            # Save the filled document
            doc.save(output_path)
            
            self.logger.info(f"Successfully filled {filled_count} placeholders in {docx_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"Error filling placeholders in {docx_path}: {str(e)}")
            return False
    
    def _fill_placeholders_in_paragraph(self, paragraph, placeholder_fills: Dict[str, str]) -> int:
        """
        Fill placeholders in a single paragraph.
        
        Args:
            paragraph: DOCX paragraph object
            placeholder_fills: Dictionary mapping placeholder text to filled content
            
        Returns:
            int: Number of placeholders filled
        """
        filled_count = 0
        
        # Combine all runs in the paragraph for processing
        full_text = ""
        run_mapping = []
        
        for run_idx, run in enumerate(paragraph.runs):
            start_pos = len(full_text)
            full_text += run.text
            end_pos = len(full_text)
            run_mapping.append({
                "run": run,
                "start": start_pos,
                "end": end_pos,
                "original_text": run.text
            })
        
        if not full_text:
            return 0
        
        # Find and replace placeholders
        modified_text = full_text
        
        for pattern in self.placeholder_patterns:
            matches = list(re.finditer(pattern, modified_text))
            for match in reversed(matches):  # Process in reverse order to maintain positions
                placeholder_text = match.group(1).strip()
                
                if placeholder_text in placeholder_fills:
                    replacement_text = placeholder_fills[placeholder_text]
                    
                    # Replace the placeholder
                    modified_text = modified_text[:match.start()] + replacement_text + modified_text[match.end():]
                    filled_count += 1
        
        # If text was modified, update the runs
        if modified_text != full_text:
            self._update_paragraph_runs(paragraph, run_mapping, full_text, modified_text)
        
        return filled_count
    
    def _update_paragraph_runs(self, paragraph, run_mapping: List[Dict], original_text: str, modified_text: str):
        """
        Update paragraph runs with modified text while preserving formatting.
        
        Args:
            paragraph: DOCX paragraph object
            run_mapping: Mapping of runs to their positions
            original_text: Original text content
            modified_text: Modified text content
        """
        try:
            # Clear existing runs
            for run in paragraph.runs:
                run.clear()
            
            # If no runs exist, create a new one
            if not paragraph.runs:
                paragraph.add_run(modified_text)
                return
            
            # Use the first run for the new content, preserving its formatting
            first_run = paragraph.runs[0]
            first_run.text = modified_text
            
            # Clear other runs
            for run in paragraph.runs[1:]:
                run.clear()
            
        except Exception as e:
            self.logger.warning(f"Error updating paragraph runs: {str(e)}")
            # Fallback: create a new run with the modified text
            paragraph.clear()
            paragraph.add_run(modified_text)
    
    def validate_docx(self, docx_path: str) -> Tuple[bool, str]:
        """
        Validate that a file is a valid DOCX template.
        
        Args:
            docx_path: Path to the file to validate
            
        Returns:
            Tuple[bool, str]: (is_valid, error_message)
        """
        try:
            doc = Document(docx_path)
            
            # Basic validation
            if len(doc.paragraphs) == 0 and len(doc.tables) == 0:
                return False, "DOCX file appears to be empty"
            
            # Check for placeholders
            placeholders = self.extract_placeholders(docx_path)
            if not placeholders:
                return False, "No placeholders found in DOCX template"
            
            return True, ""
            
        except Exception as e:
            return False, f"Invalid DOCX file: {str(e)}"
    
    def get_document_metadata(self, docx_path: str) -> Dict[str, Any]:
        """
        Extract metadata from a DOCX document.
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            Dict: Document metadata
        """
        try:
            doc = Document(docx_path)
            
            metadata = {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "file_size": Path(docx_path).stat().st_size,
                "created": doc.core_properties.created,
                "modified": doc.core_properties.modified,
                "title": doc.core_properties.title,
                "author": doc.core_properties.author,
                "subject": doc.core_properties.subject,
                "keywords": doc.core_properties.keywords
            }
            
            return metadata
            
        except Exception as e:
            self.logger.error(f"Error extracting metadata from {docx_path}: {str(e)}")
            return {}


# Convenience functions
def extract_placeholders_from_docx(docx_path: str) -> List[Placeholder]:
    """
    Convenience function to extract placeholders from a DOCX template.
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        List[Placeholder]: List of found placeholders
    """
    parser = DOCXParser()
    return parser.extract_placeholders(docx_path)


def fill_docx_template(docx_path: str, placeholder_fills: Dict[str, str], output_path: str) -> bool:
    """
    Convenience function to fill a DOCX template.
    
    Args:
        docx_path: Path to the original DOCX template
        placeholder_fills: Dictionary mapping placeholder text to filled content
        output_path: Path for the output filled document
        
    Returns:
        bool: True if successful
    """
    parser = DOCXParser()
    return parser.fill_placeholders(docx_path, placeholder_fills, output_path)


def validate_docx_template(docx_path: str) -> Tuple[bool, str]:
    """
    Convenience function to validate a DOCX template.
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        Tuple[bool, str]: (is_valid, error_message)
    """
    parser = DOCXParser()
    return parser.validate_docx(docx_path)
