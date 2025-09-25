"""
Template processing routes for uploading templates and generating filled documents.
Handles template uploads, processing, and file downloads with real-time progress updates.
"""

import os
import uuid
import logging
import asyncio
from typing import List, Optional
from pathlib import Path
from datetime import datetime
import zipfile
import tempfile

from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Depends, BackgroundTasks
from fastapi.responses import JSONResponse, FileResponse
from sqlalchemy.orm import Session
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

from models.db_models import TemplateTask, TemplateFile, create_tables
from parsers.docx_parser import DOCXParser, validate_docx_template
from services.template_filler_service import TemplateFillerService
from api.websocket_routes import (
    send_progress_update, 
    send_file_completion, 
    send_task_completion,
    send_error_notification
)
from utils.config import settings
from utils.validators import FileUploadValidator

logger = logging.getLogger(__name__)

# Database setup
engine = create_engine(settings.database_url)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)

router = APIRouter()


def get_db():
    """Get database session."""
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


@router.post("/upload")
async def upload_templates(
    background_tasks: BackgroundTasks,
    files: List[UploadFile] = File(...),
    user_context: Optional[str] = Form(None),
    process_flow_image: Optional[str] = Form(None),
    db: Session = Depends(get_db)
):
    """
    Upload template files for processing.
    
    Args:
        background_tasks: FastAPI background tasks
        files: List of uploaded template files
        user_context: Additional user context
        process_flow_image: Base64 encoded process flow image
        db: Database session
        
    Returns:
        JSONResponse: Upload result with task ID
    """
    try:
        if not files:
            raise HTTPException(status_code=400, detail="No files provided")
        
        # Validate files
        file_validator = FileUploadValidator()
        validated_files = []
        
        for file in files:
            if not file.filename:
                raise HTTPException(status_code=400, detail="One or more files have no filename")
            
            # Validate file type and size
            if not file_validator.validate_file_type(file.filename, file.content_type):
                raise HTTPException(
                    status_code=400, 
                    detail=f"Invalid file type for {file.filename}. Only DOCX files are allowed."
                )
            
            # Read and validate file content
            content = await file.read()
            if not file_validator.validate_file_size(len(content)):
                raise HTTPException(
                    status_code=400, 
                    detail=f"File {file.filename} exceeds maximum allowed size."
                )
            
            # Validate DOCX format
            is_valid, error_msg = validate_docx_from_content(content)
            if not is_valid:
                raise HTTPException(
                    status_code=400, 
                    detail=f"Invalid DOCX file {file.filename}: {error_msg}"
                )
            
            validated_files.append({
                "file": file,
                "content": content,
                "filename": file_validator.sanitize_filename(file.filename)
            })
        
        # Create task record
        task_id = str(uuid.uuid4())
        db_task = TemplateTask(
            task_id=task_id,
            status="pending",
            total_files=len(validated_files),
            user_context=user_context,
            process_flow_description=None  # Will be set after image analysis
        )
        db.add(db_task)
        db.commit()
        db.refresh(db_task)
        
        # Save files and create file records
        file_records = []
        upload_dir = Path(settings.upload_dir) / "templates"
        upload_dir.mkdir(parents=True, exist_ok=True)
        
        for i, file_data in enumerate(validated_files):
            # Save file
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            file_path = upload_dir / f"{timestamp}_{file_data['filename']}"
            
            with open(file_path, "wb") as f:
                f.write(file_data["content"])
            
            # Create file record
            db_file = TemplateFile(
                task_id=db_task.id,
                filename=file_data["filename"],
                original_filename=file_data["file"].filename,
                file_path=str(file_path)
            )
            db.add(db_file)
            file_records.append(db_file)
        
        db.commit()
        
        # Start background processing
        background_tasks.add_task(
            process_templates_background,
            task_id,
            [str(f.file_path) for f in file_records],
            user_context,
            process_flow_image
        )
        
        logger.info(f"Templates uploaded successfully: {len(validated_files)} files (Task ID: {task_id})")
        
        return JSONResponse(
            status_code=200,
            content={
                "success": True,
                "message": "Templates uploaded successfully",
                "task_id": task_id,
                "total_files": len(validated_files),
                "status": "processing"
            }
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error uploading templates: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error uploading templates: {str(e)}")


async def process_templates_background(
    task_id: str,
    file_paths: List[str],
    user_context: Optional[str],
    process_flow_image: Optional[str]
):
    """
    Background task to process templates and generate filled documents.
    
    Args:
        task_id: Task ID
        file_paths: List of template file paths
        user_context: Additional user context
        process_flow_image: Base64 encoded process flow image
    """
    db = SessionLocal()
    try:
        logger.info(f"Starting background processing for task {task_id}")
        
        # Update task status
        db_task = db.query(TemplateTask).filter(TemplateTask.task_id == task_id).first()
        if not db_task:
            logger.error(f"Task {task_id} not found")
            return
        
        db_task.status = "processing"
        db.commit()
        
        # Send initial progress update
        await send_progress_update(
            task_id=task_id,
            status="processing",
            message=f"Starting processing of {len(file_paths)} template files",
            progress=0
        )
        
        # Initialize template filler service
        filler_service = TemplateFillerService()
        
        # Process each file
        completed_files = 0
        failed_files = 0
        generated_files = []
        
        for i, file_path in enumerate(file_paths):
            try:
                # Update file status
                db_file = db.query(TemplateFile).filter(
                    TemplateFile.task_id == db_task.id,
                    TemplateFile.file_path == file_path
                ).first()
                
                if db_file:
                    db_file.status = "processing"
                    db_file.processing_started = datetime.now()
                    db.commit()
                
                # Send file start notification
                await send_progress_update(
                    task_id=task_id,
                    status="processing_file",
                    message=f"Processing file {i+1}/{len(file_paths)}",
                    progress=int((i / len(file_paths)) * 80),
                    file_index=i,
                    current_file=Path(file_path).name
                )
                
                # Create progress callback
                async def progress_callback(progress_data):
                    await send_progress_update(
                        task_id=task_id,
                        status=progress_data.get("status", "processing"),
                        message=progress_data.get("message", ""),
                        progress=int((i / len(file_paths)) * 80 + (progress_data.get("progress", 0) * 0.8)),
                        file_index=i,
                        current_file=Path(file_path).name
                    )
                
                # Process the template
                result = await filler_service.process_template_file(
                    template_path=file_path,
                    user_context=user_context,
                    process_flow_description=None,  # Will be handled by the service
                    process_flow_image=process_flow_image,
                    progress_callback=progress_callback
                )
                
                if result["success"]:
                    # Update file status
                    if db_file:
                        db_file.status = "completed"
                        db_file.processing_completed = datetime.now()
                        db_file.generated_file_path = result["output_path"]
                        db_file.placeholder_count = result["placeholders_filled"]
                        db.commit()
                    
                    completed_files += 1
                    generated_files.append(result["output_path"])
                    
                    # Send file completion notification
                    await send_file_completion(
                        task_id=task_id,
                        file_index=i,
                        filename=result["output_filename"],
                        completed_files=completed_files,
                        total_files=len(file_paths),
                        download_url=f"/api/template/download/{db_file.id if db_file else i}"
                    )
                    
                    logger.info(f"Successfully processed file {i+1}/{len(file_paths)}: {result['output_filename']}")
                else:
                    # Update file status
                    if db_file:
                        db_file.status = "failed"
                        db_file.error_message = "Processing failed"
                        db.commit()
                    
                    failed_files += 1
                    logger.error(f"Failed to process file {i+1}/{len(file_paths)}: {file_path}")
                
            except Exception as e:
                logger.error(f"Error processing file {file_path}: {str(e)}")
                
                # Update file status
                if db_file:
                    db_file.status = "failed"
                    db_file.error_message = str(e)
                    db.commit()
                
                failed_files += 1
                
                # Send error notification
                await send_error_notification(
                    task_id=task_id,
                    error_message=f"Error processing file {Path(file_path).name}: {str(e)}"
                )
        
        # Update task status
        db_task.status = "completed" if completed_files > 0 else "failed"
        db_task.completed_files = completed_files
        db_task.failed_files = failed_files
        db.commit()
        
        # Send final completion notification
        if completed_files > 0:
            await send_task_completion(
                task_id=task_id,
                success=True,
                message=f"Successfully processed {completed_files}/{len(file_paths)} files",
                completed_files=completed_files,
                failed_files=failed_files,
                generated_files=generated_files
            )
        else:
            await send_task_completion(
                task_id=task_id,
                success=False,
                message="All files failed to process",
                completed_files=completed_files,
                failed_files=failed_files
            )
        
        logger.info(f"Completed processing task {task_id}: {completed_files} successful, {failed_files} failed")
        
    except Exception as e:
        logger.error(f"Error in background processing for task {task_id}: {str(e)}")
        
        # Update task status to failed
        try:
            db_task = db.query(TemplateTask).filter(TemplateTask.task_id == task_id).first()
            if db_task:
                db_task.status = "failed"
                db.commit()
        except Exception as db_error:
            logger.error(f"Error updating task status: {str(db_error)}")
        
        # Send error notification
        await send_error_notification(
            task_id=task_id,
            error_message=f"Task processing failed: {str(e)}"
        )
    
    finally:
        db.close()


def validate_docx_from_content(content: bytes) -> tuple[bool, str]:
    """
    Validate DOCX file from content bytes.
    
    Args:
        content: DOCX file content as bytes
        
    Returns:
        tuple: (is_valid, error_message)
    """
    try:
        import tempfile
        from docx import Document
        
        # Write content to temporary file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_file.write(content)
            temp_file.flush()
            
            # Validate using python-docx
            doc = Document(temp_file.name)
            
            if len(doc.paragraphs) == 0 and len(doc.tables) == 0:
                doc.close()
                os.unlink(temp_file.name)
                return False, "DOCX file appears to be empty"
            
            doc.close()
            os.unlink(temp_file.name)
            return True, ""
            
    except Exception as e:
        return False, f"Invalid DOCX file: {str(e)}"


@router.get("/task/{task_id}/status")
async def get_task_status(task_id: str, db: Session = Depends(get_db)):
    """
    Get processing status for a task.
    
    Args:
        task_id: Task ID
        db: Database session
        
    Returns:
        JSONResponse: Task status
    """
    try:
        task = db.query(TemplateTask).filter(TemplateTask.task_id == task_id).first()
        
        if not task:
            raise HTTPException(status_code=404, detail="Task not found")
        
        # Get file details
        files = db.query(TemplateFile).filter(TemplateFile.task_id == task.id).all()
        
        file_details = []
        for file in files:
            file_details.append({
                "id": file.id,
                "filename": file.filename,
                "original_filename": file.original_filename,
                "status": file.status,
                "error_message": file.error_message,
                "generated_file_path": file.generated_file_path,
                "placeholder_count": file.placeholder_count,
                "processing_started": file.processing_started.isoformat() if file.processing_started else None,
                "processing_completed": file.processing_completed.isoformat() if file.processing_completed else None
            })
        
        return JSONResponse(
            status_code=200,
            content={
                "success": True,
                "task_id": task.task_id,
                "status": task.status,
                "total_files": task.total_files,
                "completed_files": task.completed_files,
                "failed_files": task.failed_files,
                "user_context": task.user_context,
                "process_flow_description": task.process_flow_description,
                "created_at": task.created_at.isoformat(),
                "updated_at": task.updated_at.isoformat(),
                "files": file_details
            }
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting task status: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error getting task status: {str(e)}")


@router.get("/download/{file_id}")
async def download_generated_file(file_id: int, db: Session = Depends(get_db)):
    """
    Download a single generated file.
    
    Args:
        file_id: File ID
        db: Database session
        
    Returns:
        FileResponse: Generated file
    """
    try:
        file_record = db.query(TemplateFile).filter(TemplateFile.id == file_id).first()
        
        if not file_record:
            raise HTTPException(status_code=404, detail="File not found")
        
        if file_record.status != "completed" or not file_record.generated_file_path:
            raise HTTPException(status_code=400, detail="File not ready for download")
        
        if not os.path.exists(file_record.generated_file_path):
            raise HTTPException(status_code=404, detail="Generated file not found on disk")
        
        return FileResponse(
            path=file_record.generated_file_path,
            filename=file_record.filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error downloading file: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error downloading file: {str(e)}")


@router.get("/download_all/{task_id}")
async def download_all_files(task_id: str, db: Session = Depends(get_db)):
    """
    Download all generated files as a ZIP archive.
    
    Args:
        task_id: Task ID
        db: Database session
        
    Returns:
        FileResponse: ZIP archive
    """
    try:
        task = db.query(TemplateTask).filter(TemplateTask.task_id == task_id).first()
        
        if not task:
            raise HTTPException(status_code=404, detail="Task not found")
        
        if task.status != "completed":
            raise HTTPException(status_code=400, detail="Task not completed yet")
        
        # Get completed files
        files = db.query(TemplateFile).filter(
            TemplateFile.task_id == task.id,
            TemplateFile.status == "completed",
            TemplateFile.generated_file_path.isnot(None)
        ).all()
        
        if not files:
            raise HTTPException(status_code=404, detail="No completed files found")
        
        # Create ZIP file
        zip_filename = f"generated_documents_{task_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
        zip_path = os.path.join(settings.generated_dir, zip_filename)
        
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for file in files:
                if os.path.exists(file.generated_file_path):
                    # Add file to ZIP with original filename
                    zip_file.write(
                        file.generated_file_path,
                        file.filename
                    )
        
        return FileResponse(
            path=zip_path,
            filename=zip_filename,
            media_type="application/zip"
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error creating ZIP file: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error creating ZIP file: {str(e)}")


@router.get("/tasks")
async def list_tasks(
    skip: int = 0,
    limit: int = 50,
    status: Optional[str] = None,
    db: Session = Depends(get_db)
):
    """
    List processing tasks.
    
    Args:
        skip: Number of tasks to skip
        limit: Maximum number of tasks to return
        status: Filter by task status
        db: Database session
        
    Returns:
        JSONResponse: List of tasks
    """
    try:
        query = db.query(TemplateTask)
        
        if status:
            query = query.filter(TemplateTask.status == status)
        
        tasks = query.order_by(TemplateTask.created_at.desc()).offset(skip).limit(limit).all()
        
        result = []
        for task in tasks:
            result.append({
                "task_id": task.task_id,
                "status": task.status,
                "total_files": task.total_files,
                "completed_files": task.completed_files,
                "failed_files": task.failed_files,
                "created_at": task.created_at.isoformat(),
                "updated_at": task.updated_at.isoformat()
            })
        
        return JSONResponse(
            status_code=200,
            content={
                "success": True,
                "tasks": result,
                "total": len(result)
            }
        )
        
    except Exception as e:
        logger.error(f"Error listing tasks: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error listing tasks: {str(e)}")


@router.delete("/task/{task_id}")
async def delete_task(task_id: str, db: Session = Depends(get_db)):
    """
    Delete a processing task and its associated files.
    
    Args:
        task_id: Task ID
        db: Database session
        
    Returns:
        JSONResponse: Deletion result
    """
    try:
        task = db.query(TemplateTask).filter(TemplateTask.task_id == task_id).first()
        
        if not task:
            raise HTTPException(status_code=404, detail="Task not found")
        
        # Get associated files
        files = db.query(TemplateFile).filter(TemplateFile.task_id == task.id).all()
        
        # Delete files from disk
        for file in files:
            try:
                if file.file_path and os.path.exists(file.file_path):
                    os.remove(file.file_path)
                if file.generated_file_path and os.path.exists(file.generated_file_path):
                    os.remove(file.generated_file_path)
            except Exception as e:
                logger.warning(f"Could not delete file: {str(e)}")
        
        # Delete file records (cascade should handle this)
        db.query(TemplateFile).filter(TemplateFile.task_id == task.id).delete()
        
        # Delete task
        db.delete(task)
        db.commit()
        
        logger.info(f"Successfully deleted task {task_id}")
        
        return JSONResponse(
            status_code=200,
            content={
                "success": True,
                "message": f"Task {task_id} deleted successfully"
            }
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting task: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error deleting task: {str(e)}")
